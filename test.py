import os
import time
from typing import Dict, List
from textwrap import dedent
from functools import lru_cache
from io import BytesIO
from docx import Document
from docx.shared import Inches
import re
import streamlit as st
from langgraph.graph import StateGraph, END
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.runnables import RunnableLambda
from langchain_openai import ChatOpenAI
from tavily import TavilyClient

# ===================
# Agent Implementation
# ===================

class MyLangGraphAgent:
    def __init__(self, max_retries: int = 3):
        self.max_retries = max_retries
        self.tavily = TavilyClient(api_key=st.secrets["TAVILY_API_KEY"])
        self.llm = self._setup_llm()
        self.graph = self._setup_graph()

    def _setup_llm(self) -> ChatOpenAI:
        return ChatOpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=st.secrets["OPENAI_API_KEY"],
            model="openai/gpt-4o",
            temperature=0.3,
            max_tokens=2000,
            timeout=30,
            max_retries=3
        )

    @lru_cache(maxsize=100)
    def _search_web(self, job: str) -> str:
        """Cached search function with retry logic"""
        query = (
            f'site:forbes.com OR site:techcrunch.com OR site:mit.edu OR site:gartner.com '
            f'("AI" OR "artificial intelligence") AND ("{job}" OR "{job} industry") '
            f'AND ("use case" OR "application" OR "implementation" OR "case study") '
            f'after:2023-01-01'
        )
        
        for attempt in range(self.max_retries):
            try:
                result = self.tavily.search(
                    query=query,
                    search_depth="advanced",
                    max_results=7,
                    include_raw_content=True
                )
                if not result["results"]:
                    return "No results found."
                return "\n\n".join([
                    f"{r['title']}:\n{r['content']}\nSource: {r['url']}" 
                    for r in result["results"][:5]
                ])
            except Exception as e:
                if attempt == self.max_retries - 1:
                    return f"Search failed after {self.max_retries} attempts: {str(e)}"
                time.sleep(2 ** attempt)

    def _format_prompt(self, job: str, search_results: str) -> List[Dict[str, str]]:
        system_prompt = dedent(f"""
        You are an expert AI research analyst specializing in industry applications of AI.
        Analyze the following information about AI in {job} and extract specific examples.

        ### Required Output Format ###
        You MUST respond with ONLY this Markdown table structure:

        ```markdown
        | Task/Function | AI Technology | Implementation Details | Impact | Source |
        |---------------|---------------|------------------------|--------|--------|
        [5-8 specific examples]
        ```

        Focus on:
        - Concrete implementations (specific tools/technologies)
        - Measurable impact (include numbers when possible)
        - Verified information (must match the sources provided)
        """)
        return [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": search_results}
        ]

    def _run_agent(self, state: Dict) -> Dict:
        """Wrapper with retry logic for LLM calls"""
        for attempt in range(self.max_retries):
            try:
                messages = state["messages"]
                response = self.llm.invoke(messages)
                return {"messages": messages + [AIMessage(content=response.content)]}
            except Exception as e:
                if attempt == self.max_retries - 1:
                    raise
                time.sleep(1 * (attempt + 1))

    def _setup_graph(self):
        builder = StateGraph(dict)
        builder.add_node("agent_node", RunnableLambda(self._run_agent))
        builder.set_entry_point("agent_node")
        builder.add_edge("agent_node", END)
        return builder.compile()

    def query(self, job: str) -> Dict:
        """Main query method with comprehensive error handling"""
        try:
            if not job or not isinstance(job, str):
                raise ValueError("Job title must be a non-empty string")
                
            search_results = self._search_web(job)
            if "failed" in search_results.lower():
                return {"error": search_results}
                
            prompt = self._format_prompt(job, search_results)
            messages = [HumanMessage(**m) for m in prompt]
            
            return self.graph.invoke(
                {"messages": messages},
                config={"recursion_limit": 50}
            )
        except Exception as e:
            return {
                "error": f"Query failed: {str(e)}",
                "suggestion": "Please try again or check your inputs"
            }

# ===================
# DOCX Conversion
# ===================

def add_hyperlink(paragraph, url, text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(
        url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def convert_to_docx(markdown_text: str) -> tuple[BytesIO, str]:
    """Convert markdown table to Word document"""
    doc = Document()
    doc.add_heading("AI Industry Insights Report", level=1)

    # Extract title from markdown if available
    title = "AI Industry Insights"
    title_match = re.search(r"AI applications in (.*?)\n", markdown_text)
    if title_match:
        title = f"AI in {title_match.group(1).strip()}"
        doc.add_heading(title, level=2)

    # Extract table data
    table_data = []
    in_table = False
    for line in markdown_text.split('\n'):
        if line.strip().startswith('|') and '|' in line:
            in_table = True
            table_data.append(line)
        elif in_table:
            break

    if table_data:
        # Process headers and rows
        headers = [h.strip() for h in table_data[0].split('|')[1:-1]]
        rows = []
        for row in table_data[2:]:
            cells = [c.strip() for c in row.split('|')[1:-1]]
            if any(cells):  # Skip empty rows
                rows.append(cells)

        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Add rows
        for row in rows:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                # Handle hyperlinks
                link_match = re.match(r"\[(.*?)\]\((https?://[^\s]+)\)", cell)
                if link_match and i == len(row)-1:  # Only for Source column
                    text, url = link_match.groups()
                    add_hyperlink(row_cells[i].paragraphs[0], url, text)
                else:
                    row_cells[i].text = cell

    # Save to bytes buffer
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Create filename
    filename = "".join(c for c in title if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')
    return output, f"{filename}.docx"

# ===================
# Streamlit UI
# ===================

def main():
    st.set_page_config(
        page_title="AI Industry Assistant",
        page_icon="üöÄ",
        layout="centered"
    )
    
    st.title("üöÄ AI in Industry Assistant")
    st.markdown("""
    Discover how AI is transforming specific professions or industries.
    Enter a job title or industry below to see real-world AI applications.
    """)
    
    with st.form("query_form"):
        job = st.text_input(
            "Enter a profession or industry:",
            placeholder="e.g., Digital Marketing, Healthcare, Finance",
            help="Be as specific as possible for better results"
        )
        submitted = st.form_submit_button("Analyze")
    
    if submitted and job:
        with st.spinner(f"üîç Researching AI applications in {job}..."):
            try:
                agent = MyLangGraphAgent()
                result = agent.query(job)
                
                if "error" in result:
                    st.error(result["error"])
                    if "suggestion" in result:
                        st.info(result["suggestion"])
                else:
                    markdown_output = result["messages"][-1].content
                    st.success("Analysis complete!")
                    
                    # Display results
                    st.subheader(f"AI Applications in {job}")
                    st.markdown(markdown_output)
                    
                    # Download option
                    docx_file, filename = convert_to_docx(markdown_output)
                    st.download_button(
                        label="üìÑ Download Report",
                        data=docx_file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
            except Exception as e:
                st.error(f"An unexpected error occurred: {str(e)}")
                st.info("Please try again later or contact support")

if __name__ == "__main__":
    main()